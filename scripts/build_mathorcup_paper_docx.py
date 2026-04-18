"""
生成 MathorCup C 题正式论文 docx（中文），数值与路径严格对齐 outputs/run_20260417_211020/。
需要: pip install python-docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
RUN = ROOT / "outputs" / "run_20260417_211020"
FIG = RUN / "all_figures"
# 使用 ASCII 文件名，避免部分 Windows 环境下控制台编码导致路径乱码
OUT = ROOT / "MathorCup_C_paper_final.docx"


def _p(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "宋体")
    run.font.size = Pt(12)
    if bold:
        run.bold = True


def _heading(doc: Document, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "黑体"
        run._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "黑体")


def _add_figure(doc: Document, path: Path, caption: str, width_cm: float = 16.0) -> None:
    if not path.is_file():
        _p(doc, f"[图缺失：{path.name}]", bold=True)
        return
    doc.add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.font.size = Pt(10.5)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "宋体")


def build() -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17)
    sec.right_margin = Cm(3.17)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run(
        "中老年人群高血脂症风险预警及干预方案优化的\n"
        "多阶段统计建模研究"
    )
    tr.bold = True
    tr.font.size = Pt(18)
    tr.font.name = "黑体"
    tr._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "黑体")

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = st.add_run("（2026 年 MathorCup 数学应用挑战赛 C 题）")
    sr.font.size = Pt(12)
    sr.font.name = "宋体"
    sr._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "宋体")

    _heading(doc, "摘  要", 1)
    abstract = (
        "人口老龄化背景下，中老年人群高血脂症防控需要融合中医体质偏颇、活动功能状态与血常规代谢信息。"
        "本文基于附件 1 给出的 1000 例样本，构建“数据治理—三视角潜状态建模—三级风险预警—痰湿患者六个月离散优化”的闭环体系。"
        "问题 1：将九种体质积分、活动量表与代谢偏离指标分别压缩为体质、活动、代谢一阶潜因子，并在其上进行二阶 PCA 得到综合潜状态"
        "（latent_state_h，PC1 方差解释率约 33.94%），结合载荷/bootstrap 论证结构稳健性，并量化九种体质在体质因子中的绝对贡献份额差异。"
        "问题 2：以锚定样本训练不含血脂诊断泄露的前置 Logistic 预警模型（L2、类别权重平衡、5 折交叉验证、Isotonic 概率校准），"
        "在连续风险分数上进行阈值网格搜索并结合 bootstrap 给出三级分层；交叉验证 ROC-AUC 约 0.896、PR-AUC 约 0.985，期望校准误差约 0.0085；"
        "全样本真实诊断口径下 AUC 约 0.922，显著高于“宽松含血脂特征”的对照；消融显示去除代谢模块后 AUC 大幅降至约 0.600。"
        "问题 3：在体质标签为 5 的痰湿患者（n=278）上构建三阶段（各 2 个月）离散决策空间，结合年龄—活动强度可行域、痰湿调理分级、"
        "训练频次、成本与平滑约束，进行情景 min-max 鲁棒枚举求解，并扫描预算上限得到帕累托式平均成效前沿；"
        "相对最小成本/最小负担启发式基线，优化方案使平均末期痰湿积分显著降低（配对检验 p 值极小）。"
        "全文输出均来自 `scripts/run_full_pipeline.py` 可追溯结果文件。"
    )
    _p(doc, abstract)

    _heading(doc, "关键词", 1)
    _p(
        doc,
        "高血脂症；中医体质；潜因子分析；Logistic 回归；概率校准；离散优化；帕累托前沿；防数据泄露",
    )

    _heading(doc, "一、问题重述", 1)
    _p(
        doc,
        "本题要求在中老年人群中，利用中医体质积分、活动量表与血常规等信息，完成三方面工作："
        "（1）筛选能表征痰湿体质严重度并预警高血脂发病风险的关键指标，分析九种体质贡献差异；"
        "（2）融合多维度特征构建低/中/高三级风险预警模型，阐明阈值选取依据，识别痰湿高风险人群的核心特征组合；"
        "（3）对体质标签为 5 的痰湿患者，在调理分级、活动强度、频次、耐受与经济成本约束下给出 6 个月个体化干预方案，"
        "总结“患者特征—最优方案”匹配规律，并给出样本 ID 为 1、2、3 的最优方案。",
    )

    _heading(doc, "二、模型假设与符号说明", 1)
    _p(
        doc,
        "1）样本为横截面观测，附件 1 已作为本题给定总体；2）血脂诊断标签用于锚点划分与事后评价，"
        "不直接进入“前置预警”主模型特征，以避免标签泄露；3）题面给出的年龄—强度可行域、调理分级成本、"
        "强度/频次对痰湿积分下降的简化规则在优化阶段按代码实现采纳；4）6 个月按 24 周计。",
    )
    _p(
        doc,
        "符号：样本 i；九种体质向量 c_i；活动向量 a_i；代谢偏离向量 m_i；"
        "一阶因子 F_i^(c),F_i^(a),F_i^(m)；二阶综合潜状态 H_i=latent_state_h；"
        "预警概率 p̂_i，连续风险分 R_i=100p̂_i；三级阈值 t_1,t_2；"
        "第 k 阶段决策 u_ik=(z_ik,q_ik,f_ik)（调理等级、活动强度、周训练次数）。",
    )

    _heading(doc, "三、数据治理与特征构造", 1)
    _p(
        doc,
        "阶段 `run_stage_01_data()` 将原始 Excel 规范化为 `governance/canonical_dataset.csv`，"
        "并基于临床阈值将 TC/TG/LDL-C/HDL-C、血糖、尿酸、BMI 等转为偏离量，"
        "构造活动风险、体质痰湿占优度、背景风险及痰湿×低活动、痰湿×BMI 偏离、代谢×低活动等交互项，"
        "保证后续各阶段共享同一主表。（详见 `argument_Guidance.md` 中公式与字段名。）",
    )

    _heading(doc, "四、问题一的模型与结果", 1)
    _p(
        doc,
        "4.1 三视角一阶潜因子：对体质 9 维、活动相关变量、代谢偏离变量分别做因子分析提取单因子，"
        "得到 constitution_factor、activity_factor（方向按实现取反解释）、metabolic_factor。"
        "载荷见于 `latent/latent_loadings.csv`，视角诊断见 `latent/latent_view_diagnostics.csv`。",
    )
    _p(
        doc,
        "4.2 二阶综合潜状态：对三类一阶因子向量进行标准化后做 PCA 一维压缩，再 MinMax 到 0—100，"
        "得到 latent_state_h。`latent/latent_second_order.csv` 中 PC1 组合系数与解释率分别为："
        "constitution_factor≈0.7001，activity_factor≈-0.3192，metabolic_factor≈0.6387，"
        "explained_variance_ratio≈0.3394。说明二阶综合更像“排序/解释指数”，不宜夸大为唯一真实潜变量。",
    )
    _p(
        doc,
        "4.3 九体质贡献差异：以体质因子载荷绝对占比衡量贡献，`constitution_contributions_to_latent.csv` 显示"
        "平和质绝对份额最高（约 0.445），痰湿质份额约 0.112，其余体质亦占一定份额，体现多偏颇共同承载体质因子。"
        "单变量与标签关联可参照 `constitution_univariate_risk_association.csv`。",
    )
    _p(doc, "4.4 稳健性：载荷与得分的 bootstrap 汇总见 `latent/latent_bootstrap_summary.csv`、`latent/latent_score_stability_summary.csv`。")

    for fn, cap in [
        ("latent_latent_loading_heatmap.png", "图1 潜因子载荷热图（三视角结构）"),
        ("latent_latent_loading_stability_forest.png", "图2 载荷稳定性森林图"),
        ("latent_latent_score_stability_boxplot.png", "图3 潜状态得分稳定性（bootstrap 相关）"),
    ]:
        _add_figure(doc, FIG / fn, cap)

    _heading(doc, "五、问题二的模型与结果", 1)
    _p(
        doc,
        "5.1 防泄露预警模型：采用 anchor_front_logistic，在仅含前置特征的 15 维输入上训练（risk_input_policy=strict_prediagnostic_only），"
        "L2 惩罚 C=10、StratifiedKFold 5 折、类权重 balanced、训练后 Isotonic 校准并以 Brier 选择校准器。"
        "系数与 CV 指标分别见 `risk/risk_model_coefficients.csv`、`risk/risk_model_cv_metrics.csv`："
        "ROC-AUC≈0.8959，PR-AUC≈0.9854，Brier≈0.0681，LogLoss≈0.1997；"
        "校准诊断见 `risk/risk_model_calibration.csv`，验证摘要中期望校准误差 ECE≈0.00847。",
    )
    _p(
        doc,
        "5.2 三级分层与阈值：`risk_thresholds.json` 给出一套可用于分层的连续分数阈值（低—中≈59.40，中—高≈83.05，单位与风险分数一致）；"
        "`risk_threshold_summary.json` 在 50 次 bootstrap 下阈值均值约为 t_1≈60.80、t_2≈89.54，阈值带宽均值约 28.74。"
        "分层统计见 `risk/risk_tier_summary.csv`：低风险 n=184 平均确诊率≈0.348，中风险 n=121≈0.488，高风险 n=695≈0.964，"
        "呈现与临床标签一致的风险梯度。",
    )
    _p(
        doc,
        "5.3 题面桥梁与结构解释：`validation/problem_bridge_heatmap.png` 及相关 CSV 说明一阶因子、二阶 H 与风险输出的关系；"
        "二阶维度诊断显示 PC1 仅占约 33.9% 方差，提示 H 更适合解释层而非单独替代预警核心。",
    )
    _p(
        doc,
        "5.4 痰湿子群规则：`rules/minimal_rules.csv` 给出最小规则集，例如「尿酸偏离>0」覆盖率约 0.845、纯度约 0.988；"
        "「活动总分<60」覆盖率约 0.80。稳定性见 `rules/rule_stability.csv` 与图。",
    )
    _p(
        doc,
        "5.5 消融、基线、显著性与防泄露：`validation/risk_model_ablation.csv`、`risk_model_benchmark.csv`、"
        "`risk_model_significance.csv`、`risk_leakage_benchmark.csv`。要点："
        "去除代谢信息后全样本诊断 AUC 约 0.600；宽松含血脂诊断特征在训练 CV 上极高，但在全样本真实诊断口径 AUC 约 0.887，"
        "低于严格模型的约 0.922——说明“看起来更准”的泄露口径不公平且不利于泛化叙事。",
    )

    for fn, cap in [
        ("risk_continuous_risk_score.png", "图4 连续风险分数分布"),
        ("risk_risk_score_by_tier_boxplot.png", "图5 三级风险分层下的风险分展示"),
        ("risk_risk_threshold_heatmap.png", "图6 阈值网格目标函数热图"),
        ("risk_risk_anchor_overlay.png", "图7 风险分与锚点叠加"),
        ("rules_rule_selection_frequency.png", "图8 规则选择频率（稳定性）"),
        ("validation_problem_bridge_heatmap.png", "图9 问题一—问题二桥梁（相关热图）"),
        ("validation_risk_threshold_bootstrap_distributions.png", "图10 阈值 bootstrap 分布"),
        ("validation_risk_tier_feature_gradient.png", "图11 三级风险层特征梯度"),
    ]:
        _add_figure(doc, FIG / fn, cap)

    _heading(doc, "六、问题三的模型与结果", 1)
    _p(
        doc,
        "6.1 人群与决策变量：筛选 constitution_label=5 的痰湿患者共 278 例（`diagnostics.json`）。"
        "每阶段决策为 (z,q,f)：中医调理等级 z∈{1,2,3}，活动强度 q∈{1,2,3}，周训练次数 f∈{1,…,10}，三阶段串联并施加年龄—活动可行域、"
        "痰湿积分对 z 的约束、成本上限、相邻阶段强度平滑等（与题面附表一致，按 `src/domain/intervention_rules.py` 实现）。"
        "求解采用情景 min-max 下的三阶段路径枚举（见 `stage_05_optimize` 与优化器配置）。",
    )
    _p(
        doc,
        "6.2 预算扫描与边际递减：`optimization/pareto_frontier_summary.csv`，在可行方案数保持 157 的前提下，"
        "预算上限从 500 元增至 2000 元，平均末期痰湿积分从约 47.21 降至约 42.47，`pareto_budget_evidence.json` 记录"
        "痰湿总改善约 4.74 且存在收益递减。",
    )
    _p(
        doc,
        "6.3 与启发式基线比较：`validation/optimization_baseline_summary.csv` 显示两类启发式基线平均末期痰湿约 50.74，"
        "优化方案约 42.47；`optimization_significance.csv` 对末期痰湿与潜状态的配对改善高度显著。",
    )
    _p(
        doc,
        "6.4 患者特征—方案规律：`validation/optimization_driver_summary.csv`、`strategy_mapping_by_risk_tier_age.csv`、"
        "策略热图等总结不同年龄、风险层、活动能力下的首阶段强度分布。",
    )

    for fn, cap in [
        ("validation_optimization_strategy_heatmap.png", "图12 患者特征到首阶段强度的策略热图"),
        ("validation_optimization_budget_shift.png", "图13 预算变化下的策略与成效联动"),
        ("validation_sample_1_2_3_plan_paths.png", "图14 样本 1/2/3 的三阶段路径"),
        ("validation_workflow_overview.png", "图15 验证层闭环总览"),
        ("tech_workflow_overview.png", "图16 总体技术路线示意（英文版，变量名与代码一致）"),
    ]:
        _add_figure(doc, FIG / fn, cap)

    _heading(doc, "七、样本 1、2、3 的最优干预方案（题面要求）", 1)
    _p(
        doc,
        "下表摘自 `optimization/sample_1_2_3_plans.csv`（预算上限 2000 元，主目标为帕累托痰湿积分）。"
        "三阶段方案以 (调理等级, 活动强度, 周频次) 列表表示，对应三个阶段各一组。",
    )
    t = doc.add_table(rows=4, cols=6)
    hdr = ["样本 ID", "末期痰湿积分", "末期潜状态", "总成本(元)", "总负担", "三阶段 (z,q,f)×3"]
    for j, h in enumerate(hdr):
        t.rows[0].cells[j].text = h
    rows = [
        (
            "1",
            "36.89",
            "22.63",
            "1212",
            "10.20",
            "[(3,1,6),(3,1,6),(3,1,6)]",
        ),
        (
            "2",
            "48.94",
            "22.14",
            "1380",
            "18.00",
            "[(1,2,10),(1,2,10),(1,2,10)]",
        ),
        (
            "3",
            "40.93",
            "8.19",
            "1968",
            "18.80",
            "[(2,2,10),(2,3,8),(2,3,9)]",
        ),
    ]
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            t.rows[i].cells[j].text = val
    _p(doc, "表1 题定三位样本的优化输出摘要")

    _heading(doc, "八、模型评价与稳健性综合", 1)
    _p(
        doc,
        "综上，本项目在统计建模层面提供交叉验证性能、概率校准、阈值 bootstrap、锚点单调性与分层确诊率梯度；"
        "在因果叙事层面提供问题一二桥梁与二阶维度诊断；在可信度层面提供消融、宽松泄露对照与配对显著性；"
        "在决策层面提供预算前沿、基线对照与个体化驱动摘要。主诊断 JSON 见 `validation/diagnostics.json`。",
    )

    _heading(doc, "九、结论与展望", 1)
    _p(
        doc,
        "本文给出了可从数据治理追溯到代码实现的中老年高血脂风险预警与痰湿干预优化全链条。"
        "创新点在于：严格区分“诊断锚点”和“前置预警特征”，以概率校准与阈值不确定度支撑三级分层叙事，"
        "并以离散鲁棒优化解释个体化方案与预算效率。展望：若有个体纵向血糖血脂复测，可引入动态状态空间模型改进外推。",
    )

    _heading(doc, "参考文献（示例格式，投稿前请按模板统一编号）", 1)
    refs = [
        "[1] Hosmer D W, Lemeshow S, Sturdivant R X. Applied Logistic Regression (3rd ed.). Wiley, 2013.",
        "[2] sklearn: LogisticRegression, IsotonicRegression, StratifiedKFold — https://scikit-learn.org/",
        "[3] 王琦. 中医体质学（相关章节）. （根据队里书目补全）",
    ]
    for r in refs:
        _p(doc, r)

    _heading(doc, "附录 A  全部结果文件索引", 1)
    _p(
        doc,
        "完整 CSV/JSON/图路径说明见仓库根目录 `outputs_Index.md` 与 `docs/论文章节与结果映射.md`，"
        "本稿插图文件名均对应 `outputs/run_20260417_211020/all_figures/`。",
    )

    _heading(doc, "附录 B  本轮输出目录总览（与 pipeline 一致）", 1)
    _p(
        doc,
        "全链路入口脚本为 `scripts/run_full_pipeline.py`，统一输出根目录为 `outputs/run_20260417_211020/`。"
        "六个阶段子目录依次为：governance（数据治理）、latent（问题一与潜状态）、risk（问题二风险）、" "rules（问题二规则挖掘）、optimization（问题三优化）、" "validation（消融/基线/显著性/防泄露/桥梁/阈值解释/优化机制等全部验证物）。"
        "图件集中于 `all_figures/`，本轮含 39 个 PNG/PDF 文件；正文图 1—图 16 为其中代表性强、信息量高的子集，"
        "其余图表仍可在同一目录与 `outputs_Index.md` 中逐项引用。",
    )

    doc.save(str(OUT))
    print("Wrote", OUT)


if __name__ == "__main__":
    build()
